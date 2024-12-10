from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium import webdriver
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
import re

# Initialize the Word Document outside the loop
doc = Document()

# Function to clean unwanted text
def clean_text(text):
    # Remove [if !supportLists], [endif], and other similar patterns
    return re.sub(r'\[.*?\]', '', text).strip()

# A set to track processed texts to avoid duplicates
processed_texts = set()

n = int(input("please enter numbers of pages\n>>>"))

# Initialize WebDriver
driver = webdriver.Chrome()  # Use the appropriate WebDriver for your browser

for x in range(n):
    try:
        # Open the webpage
        driver.get(f"http://e-books.helwan.edu.eg/storage/29946/index.html#/reader/chapter/{x}")

        # Wait for the div element with the class 'WordSection2' to load
        specific_div = WebDriverWait(driver, 15).until(
            EC.presence_of_element_located((By.CLASS_NAME, "WordSection2"))
        )

        # Get the HTML content
        html_content = specific_div.get_attribute('innerHTML')

        # Parse the HTML with BeautifulSoup
        soup = BeautifulSoup(html_content, "html.parser")

        # Extract and format content
        for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
            # If it's a heading, apply a heading style
            if element.name.startswith('h'):
                heading_text = clean_text(element.get_text(separator=" ", strip=True))
                if heading_text and heading_text not in processed_texts:  # Only add if text is not empty
                    doc.add_paragraph(heading_text, style='Heading ' + element.name[1])
                    processed_texts.add(heading_text)  # Mark as processed

            # If it's a paragraph, create a paragraph and add text
            elif element.name == 'p':
                paragraph = doc.add_paragraph()
                paragraph_text = ""  # To track text in the current paragraph

                for child in element.descendants:
                    if "</" not in child:
                        if child.string and child.string.strip():  # Check if it's a valid text node
                            text = clean_text(child.string.strip())
                            if not text or text in processed_texts or text in paragraph_text:  # Skip duplicates
                                continue

                            # Prepare styles to apply
                            styles = {
                                "bold": False,
                                "italic": False,
                                "highlight": False,
                                "small": False,
                                "strike": False,
                                "underline": False,
                                "superscript": False,
                            }

                            # Collect styles from parent tags
                            for parent in child.parents:
                                if parent.name in ['b', 'strong']:
                                    styles["bold"] = True
                                if parent.name in ['i', 'em']:
                                    styles["italic"] = True
                                if parent.name == 'mark':
                                    styles["highlight"] = True
                                if parent.name == 'small':
                                    styles["small"] = True
                                if parent.name == 'del':
                                    styles["strike"] = True
                                if parent.name in ['u', 'ins']:
                                    styles["underline"] = True
                                if parent.name == 'sup':
                                    styles["superscript"] = True

                            # Create a single run and apply all styles
                            run = paragraph.add_run(text + " ")
                            run.bold = styles["bold"]
                            run.italic = styles["italic"]
                            if styles["highlight"]:
                                run.font.highlight_color = 7  # Highlight yellow
                            if styles["small"]:
                                run.font.size = Pt(8)  # Smaller font
                            run.font.strike = styles["strike"]
                            run.underline = styles["underline"]
                            run.font.superscript = styles["superscript"]

                            # Add text to the current paragraph tracker and global tracker
                            paragraph_text += text
                            processed_texts.add(text)

                            # Debugging: print the text and applied styles
                            print(f"Text: {text}")
                            print(f"Styles Applied: {styles}")

    except Exception as e:
        print(f"An error occurred on page {x}: {e}")

# Save the Word document after the loop
doc.save("cleaned_extracted_content_with_headings.docx")

# Close the browser
driver.quit()
